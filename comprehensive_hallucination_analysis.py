"""
Comprehensive Hallucination Analysis Report Generator
Analyzes patterns in hallucination-prone questions and provides recommendations
"""

import pandas as pd
import json
import re
from collections import Counter, defaultdict
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class HallucinationPatternAnalyzer:
    """Analyze patterns in questions that lead to hallucinations"""
    
    def __init__(self):
        self.hallucination_patterns = defaultdict(list)
        self.question_types = defaultdict(list)
        self.keywords_analysis = defaultdict(list)
        
    def analyze_question_type(self, question: str) -> str:
        """Categorize question type based on content"""
        q_lower = question.lower()
        
        # Question word patterns
        if any(word in q_lower for word in ["ai là", "who is", "who was"]):
            return "Nhân vật/Tên riêng"
        elif any(word in q_lower for word in ["năm nào", "when", "khi nào"]):
            return "Thời gian"
        elif any(word in q_lower for word in ["ở đâu", "where", "tại đâu"]):
            return "Địa điểm"
        elif any(word in q_lower for word in ["bao nhiêu", "how much", "how many", "số lượng"]):
            return "Số liệu"
        elif any(word in q_lower for word in ["công thức", "formula", "phương trình"]):
            return "Công thức khoa học"
        elif any(word in q_lower for word in ["nguyên tố", "element", "hóa học"]):
            return "Hóa học"
        elif any(word in q_lower for word in ["protein", "enzyme", "gen", "dna", "rna"]):
            return "Sinh học"
        elif any(word in q_lower for word in ["hành tinh", "planet", "sao", "thiên văn"]):
            return "Thiên văn"
        elif any(word in q_lower for word in ["tốc độ", "speed", "vận tốc", "gia tốc"]):
            return "Vật lý"
        elif any(word in q_lower for word in ["gì là", "what is", "định nghĩa"]):
            return "Định nghĩa"
        else:
            return "Tổng quát"
    
    def extract_keywords(self, question: str) -> list:
        """Extract key terms from question"""
        # Remove common Vietnamese question words
        stop_words = ["là", "gì", "ai", "ở", "đâu", "bao", "nhiêu", "năm", "nào", "có", "của", "trong", "được", "và", "với"]
        
        # Clean and tokenize
        words = re.findall(r'\b\w+\b', question.lower())
        keywords = [word for word in words if len(word) > 2 and word not in stop_words]
        
        return keywords
    
    def analyze_all_results(self, results_data: dict):
        """Analyze hallucination patterns across all APIs"""
        all_hallucinations = []
        
        for (api, dataset), result in results_data.items():
            if "graded_data" in result:
                df = result["graded_data"]
                
                # Find hallucination cases
                direct_hallu = df[df['direct_hallucination'] == True]
                selfcrit_hallu = df[df['selfcrit_hallucination'] == True]
                
                for _, row in direct_hallu.iterrows():
                    all_hallucinations.append({
                        "api": api,
                        "dataset": dataset,
                        "question": row['question'],
                        "correct_answer": row['gold_answer'],
                        "llm_answer": row['direct_answer'],
                        "prompt_type": "direct"
                    })
                
                for _, row in selfcrit_hallu.iterrows():
                    all_hallucinations.append({
                        "api": api,
                        "dataset": dataset,
                        "question": row['question'],
                        "correct_answer": row['gold_answer'],
                        "llm_answer": row.get('selfcrit_final_span', row['selfcrit_answer']),
                        "prompt_type": "selfcrit"
                    })
        
        # Analyze patterns
        for case in all_hallucinations:
            question = case['question']
            question_type = self.analyze_question_type(question)
            keywords = self.extract_keywords(question)
            
            self.hallucination_patterns[question_type].append(case)
            self.question_types[question_type].append(question)
            self.keywords_analysis[question_type].extend(keywords)
    
    def get_pattern_statistics(self):
        """Get statistics about hallucination patterns"""
        stats = {}
        
        for q_type, cases in self.hallucination_patterns.items():
            api_counts = Counter([case['api'] for case in cases])
            prompt_counts = Counter([case['prompt_type'] for case in cases])
            
            # Most common keywords for this question type
            keyword_counts = Counter(self.keywords_analysis[q_type])
            common_keywords = keyword_counts.most_common(5)
            
            stats[q_type] = {
                "total_hallucinations": len(cases),
                "api_distribution": dict(api_counts),
                "prompt_distribution": dict(prompt_counts),
                "common_keywords": [word for word, count in common_keywords],
                "sample_questions": [case['question'] for case in cases[:3]]
            }
        
        return stats
    
    def generate_recommendations(self, stats: dict) -> dict:
        """Generate specific recommendations for each question type"""
        recommendations = {}
        
        for q_type, data in stats.items():
            total = data['total_hallucinations']
            recommendations[q_type] = {
                "risk_level": "Cao" if total > 5 else "Trung bình" if total > 2 else "Thấp",
                "total_cases": total,
                "recommendations": []
            }
            
            # Specific recommendations based on question type
            if q_type == "Nhân vật/Tên riêng":
                recommendations[q_type]["recommendations"] = [
                    "Thêm cảnh báo: 'Nếu không chắc chắn về tên người, hãy nói rõ là không biết'",
                    "Sử dụng prompt: 'Chỉ trả lời nếu bạn hoàn toàn chắc chắn về tên người'",
                    "Thêm context: 'Kiểm tra kỹ tên riêng trước khi trả lời'"
                ]
            elif q_type == "Thời gian":
                recommendations[q_type]["recommendations"] = [
                    "Thêm prompt: 'Nếu không chắc về năm/thời gian chính xác, hãy nói khoảng thời gian'",
                    "Cảnh báo: 'Các thông tin thời gian cần được kiểm chứng cẩn thận'",
                    "Sử dụng: 'Trả lời dạng khoảng thời gian thay vì năm cụ thể nếu không chắc'"
                ]
            elif q_type == "Số liệu":
                recommendations[q_type]["recommendations"] = [
                    "Thêm prompt: 'Chỉ đưa ra con số nếu hoàn toàn chắc chắn'",
                    "Sử dụng: 'Đưa ra khoảng số thay vì con số chính xác nếu không chắc'",
                    "Cảnh báo: 'Các con số cần được xác minh từ nguồn đáng tin cậy'"
                ]
            elif q_type == "Công thức khoa học":
                recommendations[q_type]["recommendations"] = [
                    "Thêm: 'Kiểm tra lại công thức trước khi đưa ra'",
                    "Sử dụng: 'Nếu không chắc về công thức chính xác, hãy mô tả nguyên lý thay thế'",
                    "Cảnh báo: 'Công thức khoa học cần độ chính xác cao'"
                ]
            else:
                recommendations[q_type]["recommendations"] = [
                    "Sử dụng prompt tổng quát: 'Nếu không chắc chắn, hãy thể hiện sự không chắc chắn'",
                    "Thêm: 'Kiểm tra lại thông tin trước khi trả lời'",
                    "Khuyến khích: 'Sử dụng ngôn ngữ thận trọng khi không chắc chắn'"
                ]
        
        return recommendations

def generate_comprehensive_hallucination_report(output_path: str = "comprehensive_hallucination_analysis.docx"):
    """Generate comprehensive hallucination analysis report"""
    
    # Load all results
    from ui.components.enhanced_analytics import load_all_existing_results
    
    results_data = load_all_existing_results()
    
    if not results_data:
        print("❌ No results found to analyze")
        return None
    
    print("🔍 Analyzing hallucination patterns...")
    
    # Initialize analyzer
    analyzer = HallucinationPatternAnalyzer()
    analyzer.analyze_all_results(results_data)
    
    # Get statistics and recommendations
    stats = analyzer.get_pattern_statistics()
    recommendations = analyzer.generate_recommendations(stats)
    
    # Create Word document
    doc = Document()
    
    # Title
    title = doc.add_heading("BÁO CÁO PHÂN TÍCH HALLUCINATION TỔNG HỢP", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(f"Phân tích mẫu hình hallucination và đề xuất cải thiện prompt")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Ngày tạo: {pd.Timestamp.now().strftime('%d/%m/%Y %H:%M')}")
    
    # Executive Summary
    doc.add_heading("1. TÓM TẮT ĐIỀU HÀNH", level=1)
    
    total_apis = len(set([api for api, dataset in results_data.keys()]))
    total_datasets = len(set([dataset for api, dataset in results_data.keys()]))
    total_hallucinations = sum([data['total_hallucinations'] for data in stats.values()])
    
    summary_items = [
        f"Số API được phân tích: {total_apis} (OpenAI, DeepSeek, Gemini, Ollama)",
        f"Số dataset được kiểm tra: {total_datasets}",
        f"Tổng số trường hợp hallucination: {total_hallucinations}",
        f"Số loại câu hỏi được phân loại: {len(stats)}",
        f"Loại câu hỏi có nguy cơ cao nhất: {max(stats.keys(), key=lambda x: stats[x]['total_hallucinations']) if stats else 'Không có'}"
    ]
    
    for item in summary_items:
        p = doc.add_paragraph()
        p.add_run(f"• {item}").bold = True
    
    # Detailed Analysis by Question Type
    doc.add_heading("2. PHÂN TÍCH CHI TIẾT THEO LOẠI CÂU HỎI", level=1)
    
    # Sort by hallucination count
    sorted_types = sorted(stats.items(), key=lambda x: x[1]['total_hallucinations'], reverse=True)
    
    for i, (q_type, data) in enumerate(sorted_types, 1):
        doc.add_heading(f"2.{i} {q_type}", level=2)
        
        # Statistics
        p = doc.add_paragraph()
        p.add_run("Thống kê:").bold = True
        
        stats_items = [
            f"Số trường hợp hallucination: {data['total_hallucinations']}",
            f"Phân bố theo API: {', '.join([f'{api}: {count}' for api, count in data['api_distribution'].items()])}",
            f"Phân bố theo prompt: {', '.join([f'{prompt}: {count}' for prompt, count in data['prompt_distribution'].items()])}",
            f"Từ khóa phổ biến: {', '.join(data['common_keywords'])}"
        ]
        
        for stat in stats_items:
            doc.add_paragraph(f"  • {stat}")
        
        # Sample questions
        p = doc.add_paragraph()
        p.add_run("Ví dụ câu hỏi dễ gây hallucination:").bold = True
        
        for j, question in enumerate(data['sample_questions'], 1):
            doc.add_paragraph(f"  {j}. {question}")
        
        # Recommendations
        rec_data = recommendations[q_type]
        p = doc.add_paragraph()
        p.add_run(f"Mức độ rủi ro: {rec_data['risk_level']}").bold = True
        
        p = doc.add_paragraph()
        p.add_run("Khuyến nghị cải thiện prompt:").bold = True
        
        for rec in rec_data['recommendations']:
            doc.add_paragraph(f"  ✓ {rec}")
        
        doc.add_paragraph()  # Empty line
    
    # Comprehensive Recommendations
    doc.add_heading("3. KHUYẾN NGHỊ TỔNG QUÁT", level=1)
    
    general_recommendations = [
        "**Prompt Strategy cơ bản:**",
        "• Luôn yêu cầu LLM thể hiện sự không chắc chắn khi không biết",
        "• Sử dụng cụm từ 'Nếu không chắc chắn, hãy nói không biết'",
        "• Thêm self-critique cho các câu hỏi có nguy cơ cao",
        "",
        "**Theo từng loại câu hỏi:**",
        "• Câu hỏi về tên riêng: Yêu cầu xác minh nguồn",
        "• Câu hỏi về số liệu: Chấp nhận khoảng thay vì con số chính xác",
        "• Câu hỏi về thời gian: Sử dụng khoảng thời gian thay vì năm cụ thể",
        "• Câu hỏi về công thức: Yêu cầu kiểm tra lại công thức",
        "",
        "**Monitoring và đánh giá:**",
        "• Theo dõi tỷ lệ hallucination theo từng loại câu hỏi",
        "• Thường xuyên cập nhật prompt dựa trên kết quả mới",
        "• Sử dụng self-critique prompting cho các chủ đề có nguy cơ cao"
    ]
    
    for rec in general_recommendations:
        if rec.startswith("**"):
            p = doc.add_paragraph()
            p.add_run(rec.replace("**", "")).bold = True
        else:
            doc.add_paragraph(rec)
    
    # API Specific Analysis
    doc.add_heading("4. PHÂN TÍCH THEO API", level=1)
    
    api_analysis = defaultdict(lambda: defaultdict(int))
    for q_type, data in stats.items():
        for api, count in data['api_distribution'].items():
            api_analysis[api][q_type] = count
    
    for api, type_counts in api_analysis.items():
        doc.add_heading(f"4.{list(api_analysis.keys()).index(api)+1} {api.upper()}", level=2)
        
        total_api_hallu = sum(type_counts.values())
        most_problematic = max(type_counts.items(), key=lambda x: x[1]) if type_counts else ("Không có", 0)
        
        doc.add_paragraph(f"Tổng hallucination: {total_api_hallu}")
        doc.add_paragraph(f"Loại câu hỏi problematic nhất: {most_problematic[0]} ({most_problematic[1]} cases)")
        
        # API specific recommendations
        if api.lower() == "gemini":
            doc.add_paragraph("Khuyến nghị: Gemini cần prompt conservative hơn cho câu hỏi factual")
        elif api.lower() == "openai":
            doc.add_paragraph("Khuyến nghị: OpenAI có thể sử dụng self-critique hiệu quả")
        elif api.lower() == "deepseek":
            doc.add_paragraph("Khuyến nghị: DeepSeek cho kết quả tốt với self-critique prompting")
        elif api.lower() == "ollama":
            doc.add_paragraph("Khuyến nghị: Ollama cần prompt đơn giản và rõ ràng")
    
    # Improved Prompt Templates
    doc.add_heading("5. MẪU PROMPT ĐƯỢC CẢI THIỆN", level=1)
    
    prompt_templates = {
        "Conservative Direct Prompt": """
Bạn là một trợ lý AI chính xác và thận trọng. 
QUAN TRỌNG: Nếu bạn không hoàn toàn chắc chắn về câu trả lời, hãy nói "Tôi không chắc chắn về thông tin này."
Chỉ trả lời khi bạn có độ tin cậy cao.

Câu hỏi: {question}
""",
        "Enhanced Self-Critique Prompt": """
Nhiệm vụ: Trả lời câu hỏi một cách chính xác và tự kiểm tra.

Bước 1 - Trả lời nháp: Đưa ra câu trả lời ban đầu
Bước 2 - Tự kiểm tra: 
- Tôi có chắc chắn 100% về thông tin này không?
- Có khả năng tôi nhầm lẫn hoặc bịa đặt thông tin không?
- Tôi có cần thêm cảnh báo về độ tin cậy không?
Bước 3 - Trả lời cuối cùng: Đưa ra câu trả lời sau khi đã kiểm tra

Câu hỏi: {question}
""",
        "Domain-Specific Prompt (Science)": """
Bạn là chuyên gia khoa học. Khi trả lời:
- Chỉ đưa ra thông tin khoa học đã được xác minh
- Nếu không chắc về con số chính xác, đưa ra khoảng ước tính
- Nếu không chắc về công thức, mô tả nguyên lý thay thế
- Luôn thêm "cần kiểm chứng thêm" nếu không hoàn toàn chắc chắn

Câu hỏi: {question}
"""
    }
    
    for template_name, template_content in prompt_templates.items():
        doc.add_heading(template_name, level=2)
        p = doc.add_paragraph(template_content)
        p.style = 'Intense Quote'
    
    # Save document
    doc.save(output_path)
    print(f"✅ Comprehensive report generated: {output_path}")
    
    return output_path, stats, recommendations

if __name__ == "__main__":
    generate_comprehensive_hallucination_report()