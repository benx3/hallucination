"""
Evaluation and Report Generation for Hallucination Detection
Handles grading, metrics calculation, and Word report generation
"""

import pandas as pd
import json
import re
import os
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Pt

class HallucinationEvaluator:
    """Comprehensive evaluator for hallucination detection experiments"""
    
    def __init__(self):
        self.uncertainty_patterns = [
            r"không chắc", r"không rõ", r"khó nói", r"not sure", 
            r"uncertain", r"có thể", r"might be", r"maybe", r"possibly"
        ]
        
        self.normalizations = {
            "ho chi minh city": "tp hồ chí minh",
            "graphics processing unit": "gpu",
            "sodium chloride": "natri clorua", 
            "pacific ocean": "thái bình dương",
            "atlantic ocean": "đại tây dương",
            "union châu âu": "european union",
            "speed of light": "tốc độ ánh sáng"
        }
    
    def normalize(self, text: str) -> str:
        """Normalize text for comparison"""
        if not isinstance(text, str):
            return ""
        
        text = text.strip().lower()
        
        # Apply domain-specific normalizations
        for old, new in self.normalizations.items():
            text = text.replace(old, new)
        
        # Clean whitespace
        text = re.sub(r"\s+", " ", text)
        return text
    
    def contains_uncertainty(self, text: str) -> bool:
        """Check if text contains uncertainty expressions"""
        normalized = self.normalize(text)
        return any(re.search(pattern, normalized) for pattern in self.uncertainty_patterns)
    
    def check_correctness(self, answer: str, gold_answer: str) -> bool:
        """Check if answer is correct against gold standard"""
        answer_norm = self.normalize(answer)
        gold_norm = self.normalize(gold_answer)
        
        # Direct substring match
        if gold_norm in answer_norm:
            return True
        
        # Special cases for scientific notation
        if gold_norm == "3e8":
            speed_variants = [
                "3e8", "300,000,000", "300000000", "3 x 10^8", 
                "3×10^8", "3*10^8", "299792458", "3*10**8"
            ]
            return any(variant in answer_norm for variant in speed_variants)
        
        # Mathematical equivalencies
        if "%" in gold_norm:
            # Handle percentage formats
            gold_num = re.search(r"(\d+(?:\.\d+)?)", gold_norm)
            if gold_num:
                variants = [
                    f"{gold_num.group(1)}%",
                    f"{gold_num.group(1)} phần trăm",
                    f"{float(gold_num.group(1))/100}"
                ]
                return any(variant in answer_norm for variant in variants)
        
        return False
    
    def grade_responses(self, results_df: pd.DataFrame, questions_df: pd.DataFrame) -> pd.DataFrame:
        """Grade all responses and calculate metrics"""
        graded_rows = []
        
        # Detect answer column name (case-insensitive)
        answer_col = None
        for col in questions_df.columns:
            col_lower = col.lower()
            if col_lower in ["answer", "ground_truth", "correct_answer", "gold_answer", "best_answer", "best answer"]:
                answer_col = col
                break
            elif "answer" in col_lower and ("correct" in col_lower or "best" in col_lower or "gold" in col_lower):
                answer_col = col
                break
        
        if answer_col is None:
            raise ValueError(f"No answer column found in dataset. Available columns: {list(questions_df.columns)}")
        
        for _, row in results_df.iterrows():
            question_idx = row.get("idx", 0) - 1
            if question_idx >= 0 and question_idx < len(questions_df):
                gold_answer = questions_df.iloc[question_idx][answer_col]
            else:
                gold_answer = ""
            
            # Grade direct answer
            direct_correct = self.check_correctness(row["direct_answer"], gold_answer)
            direct_uncertain = self.contains_uncertainty(row["direct_answer"])
            direct_hallucination = not direct_correct and not direct_uncertain
            
            # Grade self-critique answer
            selfcrit_final = row.get("selfcrit_final_span", row["selfcrit_answer"])
            selfcrit_correct = self.check_correctness(selfcrit_final, gold_answer)
            selfcrit_uncertain = self.contains_uncertainty(selfcrit_final)
            selfcrit_hallucination = not selfcrit_correct and not selfcrit_uncertain
            
            graded_rows.append({
                **row.to_dict(),
                "gold_answer": gold_answer,
                "direct_correct": direct_correct,
                "direct_uncertain": direct_uncertain,
                "direct_hallucination": direct_hallucination,
                "selfcrit_correct": selfcrit_correct,
                "selfcrit_uncertain": selfcrit_uncertain,
                "selfcrit_hallucination": selfcrit_hallucination
            })
        
        return pd.DataFrame(graded_rows)
    
    def calculate_metrics(self, graded_df: pd.DataFrame) -> Dict:
        """Calculate comprehensive metrics"""
        total = len(graded_df)
        if total == 0:
            return {}
        
        metrics = {
            "total_questions": total,
            "direct": {
                "correct_rate": graded_df["direct_correct"].mean(),
                "uncertainty_rate": graded_df["direct_uncertain"].mean(),
                "hallucination_rate": graded_df["direct_hallucination"].mean()
            },
            "selfcrit": {
                "correct_rate": graded_df["selfcrit_correct"].mean(),
                "uncertainty_rate": graded_df["selfcrit_uncertain"].mean(),
                "hallucination_rate": graded_df["selfcrit_hallucination"].mean()
            }
        }
        
        # Calculate improvement
        metrics["improvement"] = {
            "correct_delta": metrics["selfcrit"]["correct_rate"] - metrics["direct"]["correct_rate"],
            "hallucination_delta": metrics["direct"]["hallucination_rate"] - metrics["selfcrit"]["hallucination_rate"],
            "uncertainty_delta": metrics["selfcrit"]["uncertainty_rate"] - metrics["direct"]["uncertainty_rate"]
        }
        
        return metrics
    
    def generate_word_report(self, graded_df: pd.DataFrame, metrics: Dict, output_path: str):
        """Generate comprehensive Word report"""
        doc = Document()
        
        # Title
        title = doc.add_heading("Hallucination Detection Experiment Report", 0)
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        
        summary_data = [
            f"Total Questions: {metrics['total_questions']}",
            f"Direct Correct Rate: {metrics['direct']['correct_rate']:.1%}",
            f"Self-Critique Correct Rate: {metrics['selfcrit']['correct_rate']:.1%}",
            f"Direct Hallucination Rate: {metrics['direct']['hallucination_rate']:.1%}",
            f"Self-Critique Hallucination Rate: {metrics['selfcrit']['hallucination_rate']:.1%}",
            f"Hallucination Reduction: {metrics['improvement']['hallucination_delta']:.1%}"
        ]
        
        for item in summary_data:
            doc.add_paragraph(f"• {item}")
        
        # Detailed Metrics
        doc.add_heading("Detailed Metrics", level=1)
        
        # Create comparison table
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        
        headers = ["Metric", "Direct Prompt", "Self-Critique"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        table.cell(1, 0).text = "Correct Rate"
        table.cell(1, 1).text = f"{metrics['direct']['correct_rate']:.1%}"
        table.cell(1, 2).text = f"{metrics['selfcrit']['correct_rate']:.1%}"
        
        table.cell(2, 0).text = "Uncertainty Rate"
        table.cell(2, 1).text = f"{metrics['direct']['uncertainty_rate']:.1%}"
        table.cell(2, 2).text = f"{metrics['selfcrit']['uncertainty_rate']:.1%}"
        
        table.cell(3, 0).text = "Hallucination Rate"
        table.cell(3, 1).text = f"{metrics['direct']['hallucination_rate']:.1%}"
        table.cell(3, 2).text = f"{metrics['selfcrit']['hallucination_rate']:.1%}"
        
        # Sample Responses
        doc.add_heading("Sample Responses", level=1)
        
        # Show first 5 responses
        for i in range(min(5, len(graded_df))):
            row = graded_df.iloc[i]
            doc.add_heading(f"Question {i+1}", level=2)
            doc.add_paragraph(f"Q: {row['question']}")
            doc.add_paragraph(f"Gold Answer: {row['gold_answer']}")
            doc.add_paragraph(f"Direct: {row['direct_answer']} ({'✓' if row['direct_correct'] else '✗'})")
            doc.add_paragraph(f"Self-Critique: {row.get('selfcrit_final_span', row['selfcrit_answer'])} ({'✓' if row['selfcrit_correct'] else '✗'})")
        
        # Save document
        doc.save(output_path)
        print(f"Word report saved to: {output_path}")
    
    def run_evaluation(self, questions_csv: str, results_csv: str, output_dir: str) -> Dict:
        """Run complete evaluation pipeline"""
        # Load data
        questions_df = pd.read_csv(questions_csv)
        results_df = pd.read_csv(results_csv)
        
        # Grade responses
        graded_df = self.grade_responses(results_df, questions_df)
        
        # Calculate metrics
        metrics = self.calculate_metrics(graded_df)
        
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        # Save graded results
        graded_csv = os.path.join(output_dir, "results_graded.csv")
        graded_df.to_csv(graded_csv, index=False, encoding="utf-8")
        
        # Save metrics
        metrics_json = os.path.join(output_dir, "metrics.json")
        with open(metrics_json, "w", encoding="utf-8") as f:
            json.dump(metrics, f, indent=2, ensure_ascii=False)
        
        # Generate Word report
        report_docx = os.path.join(output_dir, "experiment_report.docx")
        self.generate_word_report(graded_df, metrics, report_docx)
        
        print(f"Evaluation complete. Results saved to: {output_dir}")
        return metrics

def main():
    """Main execution function"""
    questions_csv = os.getenv("INPUT_QA", "data/scientific_facts_basic.csv")
    results_csv = os.getenv("INPUT_RAW", "data/results/openai/results_raw.csv")
    output_dir = os.getenv("OUTPUT_DIR", "data/results/openai")
    
    evaluator = HallucinationEvaluator()
    metrics = evaluator.run_evaluation(questions_csv, results_csv, output_dir)
    
    print("Metrics Summary:")
    print(f"Correct Rate (Direct): {metrics['direct']['correct_rate']:.1%}")
    print(f"Correct Rate (Self-Critique): {metrics['selfcrit']['correct_rate']:.1%}")
    print(f"Hallucination Reduction: {metrics['improvement']['hallucination_delta']:.1%}")

if __name__ == "__main__":
    main()