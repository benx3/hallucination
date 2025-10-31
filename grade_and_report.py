# grade_and_report.py
# REQUIRE: pip install pandas python-docx
import pandas as pd
from docx import Document
from docx.shared import Pt
import os
import re
import json

INPUT_QA   = os.getenv("INPUT_QA", "questions_50_hard.csv")
INPUT_RAW  = os.getenv("INPUT_RAW", "ollama2/results_raw.csv")
OUT_GRADED = os.getenv("OUT_GRADED", "ollama2/results_graded.csv")
OUT_DOCX   = os.getenv("OUT_DOCX", "ollama2/Experiment_Report.docx")
OUT_JSON   = os.getenv("OUT_JSON", "ollama2/metrics.json")

UNCERTAINTY_PATTERNS = [
    r"không chắc", r"không rõ", r"khó nói", r"not sure", r"uncertain", r"có thể",
]

def normalize(s: str) -> str:
    if not isinstance(s, str):
        return ""
    s = s.strip().lower()
    repl = {
        "ho chi minh city":"tp hồ chí minh",
        "graphics processing unit":"graphics processing unit",
        "sodium chloride":"natri clorua",
        "pacific ocean":"thái bình dương",
        "atlantic ocean":"đại tây dương",
        "union châu âu":"european union",
    }
    for k,v in repl.items():
        s = s.replace(k, v)
    s = re.sub(r"\s+", " ", s)
    return s

def contains_uncertainty(s: str) -> bool:
    s_norm = normalize(s)
    for pat in UNCERTAINTY_PATTERNS:
        if re.search(pat, s_norm):
            return True
    return False

def check_correct(answer: str, gold: str) -> bool:
    a = normalize(answer)
    g = normalize(gold)
    if g in a:
        return True
    if g == "3e8":
        return any(tok in a for tok in ["3e8", "300,000,000", "300000000", "3 x 10^8", "3×10^8", "3*10^8"])
    return False

def main():
    qa = pd.read_csv(INPUT_QA)
    raw = pd.read_csv(INPUT_RAW)

    merged = raw.merge(qa, on="question", how="left")

    records = []
    for _, row in merged.iterrows():
        d_ans = str(row["direct_answer"])
        s_ans = str(row["selfcrit_answer"])
        s_final = str(row.get("selfcrit_final_span", "")) or s_ans
        gold = str(row["ground_truth"])

        d_correct = check_correct(d_ans, gold)
        s_correct = check_correct(s_final, gold)

        d_uncertain = contains_uncertainty(d_ans)
        s_uncertain = contains_uncertainty(s_final)

        d_hallu = (not d_correct) and (not d_uncertain)
        s_hallu = (not s_correct) and (not s_uncertain)

        records.append({
            "idx": int(row["idx"]),
            "question": row["question"],
            "gold": gold,
            "direct_answer": d_ans,
            "direct_correct": int(d_correct),
            "direct_uncertain": int(d_uncertain),
            "direct_hallucination": int(d_hallu),
            "selfcrit_final": s_final,
            "selfcrit_correct": int(s_correct),
            "selfcrit_uncertain": int(s_uncertain),
            "selfcrit_hallucination": int(s_hallu),
        })

    df = pd.DataFrame(records).sort_values("idx")
    df.to_csv(OUT_GRADED, index=False, encoding="utf-8")

    n = len(df)
    acc_direct = df["direct_correct"].mean() if n else 0.0
    acc_self = df["selfcrit_correct"].mean() if n else 0.0
    hallu_direct = df["direct_hallucination"].mean() if n else 0.0
    hallu_self = df["selfcrit_hallucination"].mean() if n else 0.0

    metrics = {
        "n_questions": n,
        "accuracy_direct": round(acc_direct, 4),
        "accuracy_selfcrit": round(acc_self, 4),
        "hallu_rate_direct": round(hallu_direct, 4),
        "hallu_rate_selfcrit": round(hallu_self, 4),
        "accuracy_gain": round(acc_self - acc_direct, 4),
        "hallu_reduction": round(hallu_direct - hallu_self, 4),
    }
    with open(OUT_JSON, "w", encoding="utf-8") as f:
        json.dump(metrics, f, ensure_ascii=False, indent=2)

    doc = Document()
    title = doc.add_paragraph("BÁO CÁO KẾT QUẢ THỰC NGHIỆM — Self-Critique Prompting (Ollama)")
    title.runs[0].bold = True; title.runs[0].font.size = Pt(14)
    doc.add_paragraph(f"Số câu hỏi: {n}")
    doc.add_paragraph(f"Accuracy — Direct: {metrics['accuracy_direct']*100:.1f}%  |  Self-Critique: {metrics['accuracy_selfcrit']*100:.1f}%")
    doc.add_paragraph(f"Hallucination rate — Direct: {metrics['hallu_rate_direct']*100:.1f}%  |  Self-Critique: {metrics['hallu_rate_selfcrit']*100:.1f}%")
    doc.add_paragraph(f"Accuracy Gain: {metrics['accuracy_gain']*100:.1f}%   |   Hallucination Reduction: {metrics['hallu_reduction']*100:.1f}%")

    doc.add_paragraph(" ")
    doc.add_paragraph("Bảng rút gọn (10 mẫu đầu):").runs[0].bold = True
    subset = df.head(10)[["idx","question","gold","direct_answer","selfcrit_final","direct_correct","selfcrit_correct"]]
    table = doc.add_table(rows=1, cols=len(subset.columns))
    hdr = table.rows[0].cells
    for j, col in enumerate(subset.columns):
        hdr[j].text = str(col)
    for _, r in subset.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(subset.columns):
            cells[j].text = str(r[col])[:800]

    doc.save(OUT_DOCX)
    print(f"Đã ghi {OUT_GRADED}, {OUT_JSON}, {OUT_DOCX}")

if __name__ == "__main__":
    main()